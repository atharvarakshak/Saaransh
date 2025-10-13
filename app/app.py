import streamlit as st
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pyvis.network import Network
import faiss
import numpy as np
import tempfile
import os

# -------- Helper Functions -------- #
def extract_text(uploaded_file):
    """Extract text from PDF, DOCX, or TXT."""
    ext = uploaded_file.name.split(".")[-1].lower()
    if ext == "pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(uploaded_file)
        return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
    elif ext == "docx":
        import docx
        doc = docx.Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs if p.text)
    elif ext == "txt":
        return uploaded_file.read().decode("utf-8")
    else:
        return ""

def build_graph(similarity_matrix, chunk_texts):
    """Create an interactive graph based on similarity scores."""
    net = Network(height="600px", width="800px", bgcolor="#222222", font_color="white")

    for i, text in enumerate(chunk_texts):
        label = text[:30] + "..."  # first 30 chars
        net.add_node(i, label=label, title=text)

    # Add edges for high similarity
    threshold = 0.75
    for i in range(len(chunk_texts)):
        for j in range(i + 1, len(chunk_texts)):
            sim_score = float(similarity_matrix[i, j])  # 👈 convert to Python float
            if sim_score > threshold:
                net.add_edge(i, j, value=sim_score)

    return net


# -------- Streamlit UI -------- #
st.set_page_config(page_title="GraphRAG Document Explorer", layout="wide")
st.title("📚 GraphRAG - Multi-Document Visualization")

uploaded_files = st.file_uploader("Upload your documents", type=["pdf", "docx", "txt"], accept_multiple_files=True)

if uploaded_files:
    st.success(f"{len(uploaded_files)} files uploaded successfully!")

    # Extract & chunk text
    all_chunks = []
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    for file in uploaded_files:
        text = extract_text(file)
        chunks = splitter.split_text(text)
        all_chunks.extend(chunks)

    st.write(f"Total Chunks: {len(all_chunks)}")

    # Generate embeddings
    model = SentenceTransformer('all-MiniLM-L6-v2')
    embeddings = model.encode(all_chunks, convert_to_tensor=False)

    st.success("Embedding generated successfully!")

    # Store in FAISS
    dim = embeddings.shape[1]
    index = faiss.IndexFlatL2(dim)
    index.add(np.array(embeddings))

    st.success("Stored faiss success!")

    # Build similarity matrix
    similarity_matrix = np.inner(embeddings, embeddings)
    graph = build_graph(similarity_matrix, all_chunks)


    st.success("Build similarity success!")

    # Save graph to temp file and display
    with tempfile.NamedTemporaryFile(delete=False, suffix=".html") as tmp_fiimport streamlit as st
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pyvis.network import Network
import faiss
import numpy as np
import tempfile
import os

# -------- Helper Functions -------- #
def extract_text(uploaded_file):
    """Extract text from PDF, DOCX, or TXT."""
    ext = uploaded_file.name.split(".")[-1].lower()
    if ext == "pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(uploaded_file)
        return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
    elif ext == "docx":
        import docx
        doc = docx.Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs if p.text)
    elif ext == "txt":
        return uploaded_file.read().decode("utf-8")
    else:
        return ""

def build_graph(similarity_matrix, chunk_texts):
    """Create an interactive graph based on similarity scores."""
    net = Network(height="600px", width="800px", bgcolor="#222222", font_color="white")

    for i, text in enumerate(chunk_texts):
        label = text[:30] + "..."  # first 30 chars
        net.add_node(i, label=label, title=text)

    # Add edges for high similarity
    threshold = 0.75
    for i in range(len(chunk_texts)):
        for j in range(i + 1, len(chunk_texts)):
            sim_score = float(similarity_matrix[i, j])  # 👈 convert to Python float
            if sim_score > threshold:
                net.add_edge(i, j, value=sim_score)

    return net


# -------- Streamlit UI -------- #
st.set_page_config(page_title="GraphRAG Document Explorer", layout="wide")
st.title("📚 GraphRAG - Multi-Document Visualization")

uploaded_files = st.file_uploader("Upload your documents", type=["pdf", "docx", "txt"], accept_multiple_files=True)

if uploaded_files:
    st.success(f"{len(uploaded_files)} files uploaded successfully!")

    # Extract & chunk text
    all_chunks = []
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    for file in uploaded_files:
        text = extract_text(file)
        chunks = splitter.split_text(text)
        all_chunks.extend(chunks)

    st.write(f"Total Chunks: {len(all_chunks)}")

    # Generate embeddings
    model = SentenceTransformer('all-MiniLM-L6-v2')
    embeddings = model.encode(all_chunks, convert_to_tensor=False)

    st.success("Embedding generated successfully!")

    # Store in FAISS
    dim = embeddings.shape[1]
    index = faiss.IndexFlatL2(dim)
    index.add(np.array(embeddings))

    st.success("Stored faiss success!")

    # Build similarity matrix
    similarity_matrix = np.inner(embeddings, embeddings)
    graph = build_graph(similarity_matrix, all_chunks)


    st.success("Build similarity success!")

    # Save graph to temp file and display
    with tempfile.NamedTemporaryFile(delete=False, suffix=".html") as tmp_file:
        graph.write_html(tmp_file.name)
        html_path = tmp_file.name



    with open(html_path, "r", encoding="utf-8") as f:
        html_content = f.read()

    

    st.components.v1.html(html_content, height=600, width=1200, scrolling=True)

    # Search interface
    query = st.text_input("Ask a question about your documents:")
    if query:
        q_emb = model.encode([query])
        D, I = index.search(np.array(q_emb), k=5)
        st.subheader("Top Matching Chunks")
        for idx in I[0]:
            st.write(all_chunks[idx])
le:
        graph.write_html(tmp_file.name)
        html_path = tmp_file.name



    with open(html_path, "r", encoding="utf-8") as f:
        html_content = f.read()

    

    st.components.v1.html(html_content, height=600, width=1200, scrolling=True)

    # Search interface
    query = st.text_input("Ask a question about your documents:")
    if query:
        q_emb = model.encode([query])
        D, I = index.search(np.array(q_emb), k=5)
        st.subheader("Top Matching Chunks")
        for idx in I[0]:
            st.write(all_chunks[idx])
